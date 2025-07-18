import os
import tempfile
import subprocess
from flask import Flask, render_template, request, send_file, flash, redirect, url_for
from werkzeug.utils import secure_filename
from PIL import Image
from moviepy.editor import VideoFileClip
from docx import Document
from fpdf import FPDF
from PyPDF2 import PdfReader

app = Flask(__name__)
app.secret_key = 'sua_chave_secreta'

ALLOWED_EXTENSIONS = set([
    'pdf', 'docx',
    'mp3', 'wav',
    'mp4', 'avi', 'mov', 'mkv',
    'jpg', 'jpeg', 'png', 'gif', 'bmp', 'tiff'
])

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def run_ffmpeg(input_path, output_path, extra_args=None):
    cmd = ['ffmpeg', '-y', '-i', input_path]
    if extra_args:
        cmd.extend(extra_args)
    cmd.append(output_path)
    result = subprocess.run(cmd, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
    if result.returncode != 0:
        raise RuntimeError(f"FFmpeg error: {result.stderr.decode()}")

def convert_file(input_path, output_path, input_ext, output_ext):
    image_formats = ['jpg', 'jpeg', 'png', 'gif', 'bmp', 'tiff']
    audio_formats = ['mp3', 'wav']
    video_formats = ['mp4', 'avi', 'mov', 'mkv']

    # Imagem para PDF
    if input_ext in image_formats and output_ext == 'pdf':
        with Image.open(input_path) as img:
            if img.mode in ('RGBA', 'LA'):
                background = Image.new("RGB", img.size, (255, 255, 255))
                background.paste(img, mask=img.split()[3])
                img = background
            elif img.mode != 'RGB':
                img = img.convert('RGB')
            img.save(output_path, "PDF", resolution=100.0)
        return

    # PDF para imagem (primeira página)
    if input_ext == 'pdf' and output_ext in image_formats:
        from pdf2image import convert_from_path
        images = convert_from_path(input_path)
        images[0].save(output_path)
        return

    # Imagem para imagem via ffmpeg (para formatos suportados)
    if input_ext in image_formats and output_ext in image_formats:
        run_ffmpeg(input_path, output_path)
        return

    # Áudio para Áudio via ffmpeg
    if input_ext in audio_formats and output_ext in audio_formats:
        run_ffmpeg(input_path, output_path)
        return

    # Vídeo para Vídeo via ffmpeg
    if input_ext in video_formats and output_ext in video_formats:
        run_ffmpeg(input_path, output_path)
        return

    # Vídeo para Áudio via ffmpeg
    if input_ext in video_formats and output_ext in audio_formats:
        run_ffmpeg(input_path, output_path, ['-vn'])
        return

    # DOCX para PDF (texto simples)
    if input_ext == 'docx' and output_ext == 'pdf':
        doc = Document(input_path)
        pdf = FPDF()
        pdf.add_page()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.set_font("Arial", size=12)
        for para in doc.paragraphs:
            pdf.multi_cell(0, 10, para.text)
        pdf.output(output_path)
        return

    # PDF para DOCX (texto simples)
    if input_ext == 'pdf' and output_ext == 'docx':
        reader = PdfReader(input_path)
        doc = Document()
        for page in reader.pages:
            text = page.extract_text()
            if text:
                doc.add_paragraph(text)
        doc.save(output_path)
        return

    raise ValueError(f"Conversão de {input_ext} para {output_ext} não suportada.")

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/convert', methods=['POST'])
def convert():
    if 'file' not in request.files:
        flash('Nenhum arquivo enviado')
        return redirect(request.url)
    file = request.files['file']
    if file.filename == '':
        flash('Nenhum arquivo selecionado')
        return redirect(request.url)
    if not allowed_file(file.filename):
        flash('Formato de arquivo não permitido')
        return redirect(request.url)

    input_ext = file.filename.rsplit('.', 1)[1].lower()
    output_ext = request.form.get('output-format')
    if not output_ext:
        flash('Formato de saída não selecionado')
        return redirect(request.url)

    filename_secure = secure_filename(file.filename)
    try:
        with tempfile.TemporaryDirectory() as tmpdir:
            input_path = os.path.join(tmpdir, filename_secure)
            file.save(input_path)

            output_filename = f"converted.{output_ext}"
            output_path = os.path.join(tmpdir, output_filename)

            convert_file(input_path, output_path, input_ext, output_ext)

            return send_file(output_path, as_attachment=True, download_name=output_filename)
    except Exception as e:
        flash(f'Erro na conversão: {e}')
        return redirect(url_for('index'))

if __name__ == '__main__':
    app.run(debug=False, host='0.0.0.0', port=10000)
